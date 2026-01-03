from docx import Document
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw

# 1. 创建一个简单的图标 (Simple Icon)
icon_path = 'test_icon.png'
img = Image.new('RGB', (100, 100), color = (73, 109, 137))
d = ImageDraw.Draw(img)
# 画一个简单的十字或星星作为图标
d.line((20, 50, 80, 50), fill=(255, 255, 0), width=5)
d.line((50, 20, 50, 80), fill=(255, 255, 0), width=5)
img.save(icon_path)

# 2. 创建 Word 文档
doc = Document()

# 添加标题
doc.add_heading('YUUKA 测试文档', 0)

# 添加段落
p = doc.add_paragraph('老师，这是为您生成的测试文档。')
p.add_run(' 这是一个加粗的测试文本。').bold = True

doc.add_heading('功能演示', level=1)
doc.add_paragraph('下面展示了如何插入图标：', style='List Bullet')

# 插入刚才生成的图标
doc.add_picture(icon_path, width=Inches(1.0))

doc.add_paragraph('文档生成概率：100%', style='Normal')
doc.add_paragraph('验算结果：一切正常。', style='Normal')

# 保存文档
file_name = 'Test_Document.docx'
doc.save(file_name)
print(f'Successfully generated {file_name}')
